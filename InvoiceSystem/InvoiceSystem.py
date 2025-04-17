import os
import subprocess
import datetime as dt
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox

import docx


class InvoiceAutomation:
    def __init__(self):
        self.root = tk.Tk() 
        self.root.title("Invoice Automation")
        self.root.geometry("500x600")

        self.partner_date_label = tk.label(self.root, text="Date")
        self.partner_name_label = tk.label(self.root, text="Name")
        self.partner_add_label = tk.label(self.root, text="Address")
        self.partner_cont_label = tk.label(self.root, text="Contact Number")
        self.partner_invoiceNo_label = tk.label(self.root, text="inumber")
        self.partner_service_label = tk.label(self.root, text="Service")
        self.partner_amount_label = tk.label(self.root, text="Amount")
        self.partner_price_label = tk.label(self.root, text="Price")
        self.partner_payment_method_label = tk.label(self.root, text="Payment Method")

        self.payment_methode = {
            "SBI": {
                "Recipient": "AmanInfotech",
                "Bank": "State Bank of India",
                "AccNo": "1234567890",
                "ISFC": "SBIX0000001"
                },
            "HDFC": {
                "Recipient": "AmanInfotech",
                "Bank": "HDFC Bank",
                "AccNO": "2345678901",
                "ISFC": "HDFCX0000001"
                },
            "ICICI": {
                "Recipient": "AmanInfotech",
                "Bank": "ICICI Bank",
                "AccNo": "3456789012",
                "ISFC": "ICICX0000001"
                }
            }
        self.partner_date_entry = tk.Entry(self.root)
        self.partner_name_entry = tk.Entry(self.root)
        self.partner_add_entry = tk.Entry(self.root)
        self.partner_cont_entry = tk.Entry(self.root)
        self.partner_invoiceNo_entry = tk.Entry(self.root)
        self.partner_service_entry = tk.Entry(self.root)
        self.partner_amount_entry = tk.Entry(self.root)
        self.partner_price_entry = tk.Entry(self.root)

        self.partner_payment_method = tk.StringVar(self.root)
        self.partner_payment_method.set('SBI')

        self.payment_method_dropdown = tk.OptionMenu(self.root, self.payment_method, *self.payment_methode.keys())

        self.create_button = tk.Button(self.root, text="Create Invoice", command=self.create_invoice)

        padding_options = {"fill": "x" , "expand" : True, "padx": 5, "pady": 2}

        self.partner_date_label.pack(padding_options)
        self.partner_date_entry.pack(padding_options)

        self.partner_name_label.pack(padding_options)
        self.partner_name_entry.pack(padding_options)

        self.partner_add_label.pack(padding_options)
        self.partner_add_entry.pack(padding_options)

        self.partner_cont_label.pack(padding_options)
        self.partner_cont_entry.pack(padding_options)

        self.partner_invoiceNo_label.pack(padding_options)
        self.partner_invoiceNo_entry.pack(padding_options)

        self.partner_service_label.pack(padding_options)
        self.partner_service_entry.pack(padding_options)

        self.partner_amount_label.pack(padding_options)
        self.partner_amount_entry.pack(padding_options)

        self.partner_price_label.pack(padding_options)
        self.partner_price_entry.pack(padding_options)

        self.partner_payment_method_label.pack(padding_options)
        self.payment_method_dropdown.pack(padding_options)

        self.create_button.pack(padding_options)

        self.root.mainloop()
        

    @staticmethod
    def replace_text(paragraph, old_text, new_text):
        if old_text in paragraph.text:
           paragraph.text = paragraph.text.replace(old_text, new_text)   
        
        
    def create_invoice(self):
       try:
           doc = docx.Document("template.docx")
        
           selected_method = self.selected_payment_method.get()
           payment_details = self.payment_methods[selected_method]
        
           try:
              amount = float(self.partner_amount_entry.get())
              price = float(self.partner_price_entry.get())
              cost = amount * price
            
            
              replacements = {
                 "[Date]": dt.datetime.today().strftime('%d-%m-%Y'),
                 "[Name]": self.partner_name_entry.get(),
                 "[Address]": self.partner_add_entry.get(),
                 "[Contact Number]": self.partner_cont_entry.get(),
                 "[inumber]": self.partner_invoiceNo_entry.get(),
                 "[Service]": self.partner_service_entry.get(),
                 "[Amount]": str(amount),
                 "[Single Price]": f"${price:.2f}",
                 "[Full Price]": f"${cost:.2f}",
                 "[Recipient]": payment_details['Recipient'],
                 "[Bank]": payment_details['Bank'],
                 "[AccNo]": payment_details['AccNo'],
                 "[IFSC]": payment_details['IFSC']
              }
        
           except ValueError:
              messagebox.showerror('Error', 'Invalid amount or price')
              return 
        
           for paragraph in doc.paragraphs:
               for old_text, new_text in replacements.items():
                  self.replace_text(paragraph, old_text, new_text)
            
           for table in doc.tables:
              for row in table.rows:
                 for cell in row.cells:
                     for paragraph in cell.paragraphs:
                         for old_text, new_text in replacements.items():
                             self.replace_text(paragraph, old_text, new_text)
                            
           save_path = filedialog.asksaveasfilename(
               defaultextension='.pdf',
               filetypes=[('PDF documents', '*.pdf')]
           )
        
           if not save_path:
              return
            
           temp_docx = "temp.docx"
           doc.save(temp_docx)
        
           try:
              convert(temp_docx, save_path)
              messagebox.showinfo('Success', 'Invoice created and saved successfully!')
           except Exception as e:
               messagebox.showerror('Error', f'Failed to convert to PDF: {str(e)}')
           finally:
               if os.path.exists(temp_docx):
                    os.remove(temp_docx)
                
        except Exception as e:
            messagebox.showerror('Error', f'An error occurred: {str(e)}')

if __name__ == "__main__":
    InvoiceAutomation()
